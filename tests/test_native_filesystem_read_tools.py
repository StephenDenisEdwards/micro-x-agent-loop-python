"""Characterization tests for F2 native read-only filesystem tools
(read_file, grep, glob). No TS suite existed — these pin behaviour:
line-numbering, ADR-023 truncation markers, binary refusal, .docx,
path-policy rejection, grep modes incl. the single-line trap, glob
mtime ordering / brace expansion / dot exclusion.
"""

from __future__ import annotations

import os
import tempfile
import time
import unittest
from pathlib import Path

from micro_x_agent_loop.native_tools.filesystem.paths import PathPolicy
from micro_x_agent_loop.native_tools.filesystem.read_tools import (
    GlobTool,
    GrepTool,
    ReadFileTool,
    _resolve_rg,
)


class ReadFileTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.root = os.path.realpath(self._tmp.name)
        self.policy = PathPolicy(self.root)
        self.tool = ReadFileTool(self.policy)
        self.addCleanup(self._tmp.cleanup)

    async def test_line_numbered_output(self) -> None:
        (Path(self.root) / "a.txt").write_text("one\ntwo\nthree\n", encoding="utf-8")
        r = await self.tool.execute({"path": "a.txt"})
        self.assertFalse(r.is_error)
        self.assertIn("     1\tone", r.text)
        self.assertIn("     3\tthree", r.text)
        self.assertEqual(r.structured["total_lines"], 3)
        self.assertFalse(r.structured["truncated"])

    async def test_offset_limit_and_truncation_marker(self) -> None:
        (Path(self.root) / "b.txt").write_text(
            "\n".join(f"L{i}" for i in range(1, 11)) + "\n", encoding="utf-8"
        )
        r = await self.tool.execute({"path": "b.txt", "offset": 2, "limit": 3})
        self.assertIn("     2\tL2", r.text)
        self.assertIn("     4\tL4", r.text)
        self.assertNotIn("\t L5", r.text)
        self.assertTrue(r.structured["truncated"])
        self.assertIn("[Output truncated: showed lines 2-4 of 10", r.text)
        self.assertIn("offset=5, limit=3)]", r.text)

    async def test_empty_file(self) -> None:
        (Path(self.root) / "e.txt").write_text("", encoding="utf-8")
        r = await self.tool.execute({"path": "e.txt"})
        self.assertEqual(r.text, "(file is empty)")

    async def test_offset_past_end(self) -> None:
        (Path(self.root) / "c.txt").write_text("only\n", encoding="utf-8")
        r = await self.tool.execute({"path": "c.txt", "offset": 99})
        self.assertIn("past end of file", r.text)

    async def test_binary_refused(self) -> None:
        (Path(self.root) / "bin").write_bytes(b"abc\x00def")
        r = await self.tool.execute({"path": "bin"})
        self.assertTrue(r.is_error)
        self.assertIn("refusing to read binary file", r.text)

    async def test_path_policy_rejection(self) -> None:
        with tempfile.TemporaryDirectory() as outside:
            target = os.path.join(os.path.realpath(outside), "x.txt")
            Path(target).write_text("secret", encoding="utf-8")
            r = await self.tool.execute({"path": target})
            self.assertTrue(r.is_error)
            self.assertIn("outside the allowed roots", r.text)

    async def test_docx(self) -> None:
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")
        import docx

        d = docx.Document()
        d.add_paragraph("hello docx")
        d.add_paragraph("second para")
        p = os.path.join(self.root, "doc.docx")
        d.save(p)
        r = await self.tool.execute({"path": "doc.docx"})
        self.assertFalse(r.is_error)
        self.assertIn("hello docx", r.text)
        self.assertIn("second para", r.text)


class GrepTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        try:
            _resolve_rg()
        except RuntimeError as ex:
            self.skipTest(f"ripgrep unavailable: {ex}")
        self._tmp = tempfile.TemporaryDirectory()
        self.root = os.path.realpath(self._tmp.name)
        self.policy = PathPolicy(self.root)
        self.tool = GrepTool(self.policy)
        (Path(self.root) / "f.txt").write_text(
            "alpha\nbeta TARGET\ngamma TARGET\n", encoding="utf-8"
        )
        self.addCleanup(self._tmp.cleanup)

    async def test_content_mode(self) -> None:
        r = await self.tool.execute({"pattern": "TARGET", "output_mode": "content"})
        self.assertFalse(r.is_error)
        self.assertIn("TARGET", r.text)
        self.assertEqual(r.structured["mode"], "content")

    async def test_count_mode(self) -> None:
        r = await self.tool.execute({"pattern": "TARGET", "output_mode": "count"})
        self.assertEqual(r.structured["match_count"], 2)

    async def test_no_matches(self) -> None:
        r = await self.tool.execute({"pattern": "NOPE", "output_mode": "content"})
        self.assertEqual(r.text, "(no matches)")

    async def test_count_mode_counts_occurrences_not_lines(self) -> None:
        # Discovery from this characterization test: the TS grep uses ripgrep
        # --count-matches (occurrences), NOT --count (lines). So count mode is
        # CORRECT on a single-line file — 5 <item> on one line -> 5, not 1.
        # The "grep count returns 1" framing in the ISSUE-007 saga was wrong
        # for count mode; the real line-oriented trap is content mode (below).
        (Path(self.root) / "feed.rss").write_text(
            "<rss>" + "<item>x</item>" * 5 + "</rss>", encoding="utf-8"
        )
        r = await self.tool.execute(
            {"pattern": "<item>", "path": "feed.rss", "output_mode": "count"}
        )
        self.assertEqual(r.structured["match_count"], 5)

    async def test_content_mode_single_line_returns_whole_line(self) -> None:
        # The actual line-oriented trap: content mode on a single-line file
        # yields the entire file as one matching "line".
        (Path(self.root) / "one.rss").write_text(
            "<rss>" + "<item>x</item>" * 5 + "</rss>", encoding="utf-8"
        )
        r = await self.tool.execute(
            {"pattern": "<item>", "path": "one.rss", "output_mode": "content"}
        )
        self.assertEqual(r.structured["match_count"], 1)  # one line matched
        self.assertIn("<item>x</item><item>x</item>", r.text)  # whole line back

    async def test_path_policy_rejection(self) -> None:
        with tempfile.TemporaryDirectory() as outside:
            r = await self.tool.execute(
                {"pattern": "x", "path": os.path.realpath(outside)}
            )
            self.assertTrue(r.is_error)
            self.assertIn("outside the allowed roots", r.text)


class GlobTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.root = os.path.realpath(self._tmp.name)
        self.policy = PathPolicy(self.root)
        self.tool = GlobTool(self.policy)
        self.addCleanup(self._tmp.cleanup)

    async def test_mtime_desc_and_onlyfiles(self) -> None:
        old = Path(self.root) / "old.ts"
        new = Path(self.root) / "new.ts"
        old.write_text("a", encoding="utf-8")
        new.write_text("b", encoding="utf-8")
        # Force an explicit mtime delta instead of a wall-clock sleep so the
        # test is deterministic on every filesystem (some have 1s mtime
        # resolution and a 20ms sleep wouldn't suffice anyway).
        now = time.time()
        os.utime(old, (now - 10, now - 10))
        os.utime(new, (now, now))
        os.makedirs(os.path.join(self.root, "sub.ts"))  # dir must be excluded
        r = await self.tool.execute({"pattern": "**/*.ts"})
        lines = r.text.splitlines()
        self.assertTrue(lines[0].endswith("new.ts"))
        self.assertTrue(lines[1].endswith("old.ts"))
        self.assertFalse(any("sub.ts" in ln for ln in lines))

    async def test_brace_expansion(self) -> None:
        (Path(self.root) / "x.js").write_text("a", encoding="utf-8")
        (Path(self.root) / "y.jsx").write_text("b", encoding="utf-8")
        (Path(self.root) / "z.py").write_text("c", encoding="utf-8")
        r = await self.tool.execute({"pattern": "**/*.{js,jsx}"})
        self.assertIn("x.js", r.text)
        self.assertIn("y.jsx", r.text)
        self.assertNotIn("z.py", r.text)

    async def test_dot_excluded(self) -> None:
        os.makedirs(os.path.join(self.root, ".hidden"))
        (Path(self.root) / ".hidden" / "h.ts").write_text("a", encoding="utf-8")
        (Path(self.root) / "v.ts").write_text("b", encoding="utf-8")
        r = await self.tool.execute({"pattern": "**/*.ts"})
        self.assertIn("v.ts", r.text)
        self.assertNotIn("h.ts", r.text)

    async def test_no_matches(self) -> None:
        r = await self.tool.execute({"pattern": "**/*.nope"})
        self.assertEqual(r.text, "(no matches)")


if __name__ == "__main__":
    unittest.main()
